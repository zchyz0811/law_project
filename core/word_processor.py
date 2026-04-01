from docx import Document
from docx.shared import RGBColor
from core.template_engine import TemplateEngine, VariableOccurrence
from typing import List, Dict, Tuple


class WordProcessor:
    def __init__(self):
        self.engine = TemplateEngine()

    def extract_text(self, path: str) -> str:
        """提取Word文档的所有文本"""
        doc = Document(path)
        all_text = []

        for para in doc.paragraphs:
            all_text.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text.append(cell.text)

        return "\n".join(all_text)

    def extract_variables(self, path: str):
        """从已有模板中提取变量"""
        text = self.extract_text(path)
        return self.engine.extract_variables(text)

    def apply_replacements_by_positions(self, source_path: str, output_path: str,
                                        replacement_list: List[Tuple[int, int, str]]):
        """按位置执行替换列表

        replacement_list: [(start_pos, end_pos, new_text), ...]
            - start < end: 替换该区间文本为 new_text
            - start == end: 在该位置插入 new_text
            - new_text 为空: 删除该区间文本
        """
        doc = Document(source_path)

        # 构建段落偏移映射
        segments = []  # [(para, offset, length)]
        offset = 0

        for para in doc.paragraphs:
            text = para.text
            segments.append((para, offset, len(text)))
            offset += len(text) + 1

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = para.text
                        segments.append((para, offset, len(text)))
                        offset += len(text) + 1

        # 按段落分组，每个段落收集它的替换操作
        para_replacements: Dict[int, List[Tuple[int, int, str]]] = {}
        for start, end, new_text in replacement_list:
            for idx, (para, seg_offset, seg_len) in enumerate(segments):
                seg_end = seg_offset + seg_len
                if start >= seg_offset and end <= seg_end:
                    local_start = start - seg_offset
                    local_end = end - seg_offset
                    para_replacements.setdefault(idx, []).append(
                        (local_start, local_end, new_text)
                    )
                    break

        # 对每个段落，按位置倒序执行替换
        for idx, reps in para_replacements.items():
            para = segments[idx][0]
            full_text = "".join(run.text for run in para.runs)
            # 倒序替换，不影响前面的位置
            for local_start, local_end, new_text in sorted(reps, key=lambda r: r[0], reverse=True):
                full_text = full_text[:local_start] + new_text + full_text[local_end:]
            if para.runs:
                para.runs[0].text = full_text
                for run in para.runs[1:]:
                    run.text = ""

        doc.save(output_path)

    def create_template(self, original_path: str, output_path: str,
                        replacements: Dict[str, str]):
        """将原始文档转换为模板（全局替换，保留向后兼容）"""
        doc = Document(original_path)

        for para in doc.paragraphs:
            self._replace_in_paragraph(para, replacements)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._replace_in_paragraph(para, replacements)

        doc.save(output_path)

    def _replace_in_paragraph(self, para, replacements: Dict[str, str]):
        """在段落中替换文本为占位符，保留格式"""
        full_text = "".join(run.text for run in para.runs)
        sorted_items = sorted(replacements.items(), key=lambda x: len(x[0]),
                              reverse=True)

        new_text = full_text
        for original, placeholder in sorted_items:
            new_text = new_text.replace(original, placeholder)

        if new_text != full_text and para.runs:
            para.runs[0].text = new_text
            for run in para.runs[1:]:
                run.text = ""

    def fill_and_save(self, template_path: str, output_path: str,
                      values: Dict[str, str]):
        """填充模板并保存"""
        doc = Document(template_path)

        for para in doc.paragraphs:
            self._fill_paragraph(para, values)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._fill_paragraph(para, values)

        doc.save(output_path)

    def _fill_paragraph(self, para, values: Dict[str, str]):
        """填充段落中的占位符"""
        full_text = "".join(run.text for run in para.runs)
        new_text = self.engine.fill(full_text, values)

        if new_text != full_text and para.runs:
            para.runs[0].text = new_text
            for run in para.runs[1:]:
                run.text = ""
